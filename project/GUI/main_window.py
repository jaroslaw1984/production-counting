import customtkinter
import customtkinter as ctk
import pandas as pd
import tkinter as tk
import math
import warnings
import os
import tempfile
import re
from datetime import date, timedelta, datetime
from tkinter import filedialog, messagebox, ttk
from typing import Optional, Dict, Any
from pathlib import Path
from docx import Document
from docx.shared import Cm
from project.config.db_loader import fetch_available_machines, fetch_orders_for_machines, normalize_db_df, fetch_sap_basic_profiles
from project.config.workplace_config_provider import merge_db_and_csv_config
from project.config.count_per_loader import update_count_by_shift



# stała ścieżka do pliku konfiguracyjnego
BASE_DIR = Path(__file__).resolve().parent.parent
CONFING_PATH = BASE_DIR / "config" / "profile_config.csv"
MACHINE_CONFIG_PATH = BASE_DIR / "config" / "machine_config.csv"
SHIFTS_PER_DAY = 3

ORDER_ALIASES = [
    "zlecenie", "nr zlecenia", "zlecenie nr",
    "auftrag", "auftragsnr", "auftragsnummer",
    "order", "order id"
]

GRUNDPROFIL_ALIASES = [
    "grundprofil", "grund profil", "grund-profil",
    "podkład", "podklad", "profil podstawowy"
]


# ignore specific pandas warning about SQLAlchemy
warnings.filterwarnings("ignore", message="pandas only supports SQLAlchemy connectable*")  

# --- INTERFEJS GRAFICZNY (GUI) ---
def run_app():
    app_state: dict[str, Any] = {
        "df": None,
        "df_hydra": None,
        "hydra_path": None,
        "cfg": None,
        "machine_cfg": None,
        "table_frame": None,
        "last_report_text": "",
        "last_report_data": None,
        "last_report_kind": None,
        } # Słownik do przechowywania stanu aplikacji (np. wczytany DataFrame)

    customtkinter.set_appearance_mode("Dark")  # Tryby: "System" (domyślny), "Dark", "Light"
    customtkinter.set_default_color_theme("dark-blue")  # Motywy: "blue" (domyślny), "green", "dark-blue"
    
    root = customtkinter.CTk()
    root.title("Policz produkcję")
    root.geometry("800x500")

    def center_popup(parent, popup):
        try:
            parent.update_idletasks()
            popup.update_idletasks()
            pw = popup.winfo_width()
            ph = popup.winfo_height()
            rw = parent.winfo_width()
            rh = parent.winfo_height()
            rx = parent.winfo_rootx()
            ry = parent.winfo_rooty()
            x = rx + (rw - pw) // 2
            y = ry + (rh - ph) // 2
            popup.geometry(f"+{x}+{y}")
        except Exception:
            # best-effort centering — nie przerywamy działania aplikacji
            pass

# 1) Główna siatka: 2 kolumny (lewy panel i prawa treść)
    root.grid_columnconfigure(1, weight=1)
    root.grid_rowconfigure(0, weight=1)

    default_font = customtkinter.CTkFont(family="Segoe UI", size=14, weight="bold")

# 2) Lewy panel (np. przyciski/filtry)
    left = customtkinter.CTkFrame(root)
    left.grid(row=0, column=0, sticky="ns", padx=10, pady=10)
    left.grid_columnconfigure(0, weight=1) # tylko góra-dół
    left.grid_rowconfigure(98, weight=1)  # push przyciski do góry

# 3) Prawa część (rośnie w obie strony 
    right = customtkinter.CTkFrame(root)
    right.grid(row=0, column=1, sticky="nsew", padx=10, pady=10) # rośnie w obie strony


# 4) Wnętrze prawego panelu też robimy responsywne
    right.grid_columnconfigure(0, weight=1)
    right.grid_rowconfigure(0, weight=0)  # toolbar
    right.grid_rowconfigure(1, weight=1)  # treść (textbox / tabela)
          
 # 5) Element, który ma się rozciągać (np. Text lub Treeview)
    text = customtkinter.CTkTextbox(right)
    text.grid(row=1, column=0, sticky="nsew")
    
    # --- przyciski raportu (tworzone raz) ---
    edit_btn = customtkinter.CTkButton(
        right,
        text="Edytuj raport",
        command=lambda: edit_current_report(),
        width=140
    )
    edit_btn.place_forget()

    print_btn = customtkinter.CTkButton(
        right,
        text="Drukuj raport",
        command=lambda: print_current_report(),
        width=140
    )
    print_btn.place_forget()
    
    text.configure(state="disabled")  # na start zablokowany do edycji

    
    def _set_print_visible(visible: bool) -> None:
        kind = app_state.get("last_report_kind")

        # Drukuj: pokazuj dla "sap" ORAZ dla "db"
        can_print = visible and (kind in ("sap", "db"))

        # Edytuj: tylko dla "sap"
        can_edit = visible and (kind == "sap")

        if can_print:
            print_btn.place(relx=1.0, rely=1.0, anchor="se", x=-20, y=-20)
        else:
            print_btn.place_forget()

        if can_edit:
            edit_btn.place(relx=1.0, rely=1.0, anchor="se", x=-20, y=-60)
        else:
            edit_btn.place_forget()

            
    def _norm(text: str) -> str:
        return " ".join(str(text).replace("\xa0", " ").strip().lower().split())


    def _contains_any(cell_text: str, aliases: list[str]) -> bool:
        t = _norm(cell_text)
        return any(a in t for a in aliases)


    def detect_header_row(xlsx_path: str, max_scan_rows: int = 40) -> int:
        raw = pd.read_excel(xlsx_path, engine="openpyxl", header=None, nrows=max_scan_rows)
        for r in range(len(raw)):
            row = raw.iloc[r].astype(str).tolist()
            has_order = any(_contains_any(c, ORDER_ALIASES) for c in row)
            has_gp = any(_contains_any(c, GRUNDPROFIL_ALIASES) for c in row)
            if has_order and has_gp:
                return r
        raise ValueError("Nie wykryłem wiersza nagłówków (brak Zlecenie/Auftrag lub Grundprofil).")


    def find_column(df: pd.DataFrame, aliases: list[str]) -> str:
        for col in df.columns:
            if _contains_any(col, aliases):
                return col
        raise ValueError(f"Nie znalazłem kolumny pasującej do aliasów: {aliases}")


    def load_hydra_queue(xlsx_path: str) -> pd.DataFrame:
        header_row = detect_header_row(xlsx_path)
        df = pd.read_excel(xlsx_path, engine="openpyxl", header=header_row)

        # normalizacja kolumn
        df.columns = [" ".join(str(c).replace("\xa0", " ").strip().split()) for c in df.columns]

        order_col = find_column(df, ORDER_ALIASES)
        gp_col = find_column(df, GRUNDPROFIL_ALIASES)

        out = df[[order_col, gp_col]].copy()
        out = out.rename(columns={order_col: "order_id", gp_col: "grundprofil"})
        out["order_id"] = out["order_id"].astype("string").str.strip()
        out["grundprofil"] = out["grundprofil"].astype("string").str.strip()
        out = out[(out["order_id"] != "") & (out["grundprofil"] != "")]
        return out.reset_index(drop=True)


    def _normalize_order_id(s: str) -> str:
        s = str(s).strip()
        s = re.sub(r"\.0$", "", s)      # usuń końcówkę .0
        s = s.lstrip("0")              # usuń zera z przodu (do porównań)
        return s if s != "" else "0"


    def cut_from_order(df: pd.DataFrame, start_order_id: str) -> pd.DataFrame:
        start_norm = _normalize_order_id(start_order_id)

        # zrób kolumnę pomocniczą do porównania
        tmp = df.copy()
        tmp["_order_norm"] = tmp["order_id"].map(_normalize_order_id)

        hits = tmp.index[tmp["_order_norm"] == start_norm].tolist()
        if not hits:
            # debug pomocny: pokaż 10 pierwszych zleceń jakie program widzi
            sample = tmp["order_id"].head(10).tolist()
            raise ValueError(
                f"Nie znaleziono startowego zlecenia: {start_order_id}\n"
                f"(dla porównania: pierwsze zlecenia w pliku: {sample})"
            )

        return df.loc[hits[0]:].reset_index(drop=True)


    def build_sequence(df: pd.DataFrame) -> list[str]:
        # usuwa tylko duplikaty obok siebie (bezpiecznik)
        seq = df["grundprofil"].tolist()
        out = []
        prev = None
        for gp in seq:
            if gp == prev:
                continue
            out.append(gp)
            prev = gp
        return out
   
    def ask_report_params_popup(parent) -> dict | None:
        """
        Jeden popup: wybór LINIA z listy + startowe zlecenie.
        Zwraca: {"linia": "...", "start_order_id": "..."} albo None gdy anulowano.
        """
        popup = ctk.CTkToplevel(parent)
        popup.title("Parametry raportu")
        popup.resizable(False, False)
        popup.grab_set()
        center_popup(parent, popup)

        result: dict | None = None

        # --- pobierz listę maszyn (LINIA) z DB ---
        try:
            machines = fetch_available_machines()
        except Exception as e:
            messagebox.showerror("Błąd DB", f"Nie mogę pobrać listy maszyn z DB:\n{e}")
            popup.destroy()
            return None

        machines = sorted({m.strip() for m in machines if str(m).strip()})

        ctk.CTkLabel(popup, text="Wybierz linię (LINIA):").pack(anchor="w", padx=12, pady=(12, 4))
        linia_var = tk.StringVar(value=machines[0] if machines else "")
        linia_cb = ctk.CTkComboBox(popup, values=machines, variable=linia_var, width=260)
        linia_cb.pack(anchor="w", padx=12, pady=(0, 10))

        ctk.CTkLabel(popup, text="Startowe zlecenie nowej grupy:").pack(anchor="w", padx=12, pady=(0, 4))
        start_var = tk.StringVar(value="")
        start_entry = ctk.CTkEntry(popup, textvariable=start_var, width=260)
        start_entry.pack(anchor="w", padx=12, pady=(0, 12))
        start_entry.focus_set()

        def on_ok():
            nonlocal result
            linia = (linia_var.get() or "").strip().upper()
            start_order_id = (start_var.get() or "").strip()

            if not linia:
                messagebox.showwarning("Brak linii", "Wybierz linię.")
                return
            if not start_order_id:
                messagebox.showwarning("Brak zlecenia", "Podaj startowe zlecenie.")
                return

            result = {"linia": linia, "start_order_id": start_order_id}
            popup.destroy()

        def on_cancel():
            popup.destroy()

        btns = ctk.CTkFrame(popup, fg_color="transparent")
        btns.pack(fill="x", padx=12, pady=(0, 12))

        ctk.CTkButton(btns, text="Anuluj", command=on_cancel).pack(side="right", padx=(8, 0))
        ctk.CTkButton(btns, text="OK", command=on_ok).pack(side="right")

        popup.wait_window()
        return result
   
    
    # def ask_line_popup(parent) -> str | None:
    #     dialog = ctk.CTkInputDialog(text="Podaj linię (np. WLO-U005):", title="Linia SAP")
    #     val = dialog.get_input()
    #     if val is None:
    #         return None
    #     val = str(val).strip().upper()
    #     return val if val else None
    

    def generate_logistics_report():
        # 1) wybór pliku Hydry
        file_path = filedialog.askopenfilename(
            title="Wczytaj eksport Hydry (.xlsx)",
            filetypes=[("Excel", "*.xlsx")]
        )
        if not file_path:
            return

        try:
            df_hydra = load_hydra_queue(file_path)
        except Exception as e:
            messagebox.showerror("Błąd wczytania Hydry", str(e))
            return

        # 2) start zlecenia
        params = ask_report_params_popup(root)
        if not params:
            return

        linia_value = params["linia"]
        start_order_id = params["start_order_id"]

        try:
            df_group = cut_from_order(df_hydra, start_order_id)
        except Exception as e:
            messagebox.showerror("Błąd startu grupy", str(e))
            return

        seq = build_sequence(df_group)
 
        # 3) na MVP wypisz w oknie wynik
        report_text = "=== Kolejność podstaw (Hydra) ===\n\n" + "\n".join(
            f"{i+1}. {gp}" for i, gp in enumerate(seq)
        )

        text.configure(state="normal")
        text.delete("1.0", "end")
        text.insert("end", report_text)
        text.configure(state="disabled")
        _upadate_placeholder_visibility()        
        
        # --- SAP -> mapy ilości ---
        sap_qty = {}
        sap_jm = {}
        sap_szt = {}
        sap_user = None
        sap_date = None
        
        # linia_value = ask_line_popup(root)
        # if not linia_value:
        #     return
        
        day_value = date.today()
        
        try:
            df_sap = fetch_sap_basic_profiles(linia=linia_value, day=day_value)
        except Exception as e:
            messagebox.showerror("SAP/DB error", f"{type(e).__name__}: {e}")
            return

        if df_sap is None or df_sap.empty:
            messagebox.showwarning("Brak danych SAP", f"Brak danych dla: {linia_value} / {day_value}")
            return           

        for _, r in df_sap.iterrows():
            idx = str(r["INDEKS"]).strip()

            # --- ILOŚĆ (metry) ---
            qty = r["ILOSC"]
            if isinstance(qty, str):
                qty = qty.replace(",", ".")
            try:
                qty = float(qty)
            except Exception:
                qty = 0.0

            # --- SZTUKI ---
            szt = r.get("IL_SZT", 0)
            try:
                szt = int(szt)
            except Exception:
                szt = 0

            # --- JM ---
            jm = str(r.get("JM", "M")).strip()

            # --- sumowanie ---
            sap_qty[idx] = sap_qty.get(idx, 0.0) + qty
            sap_szt[idx] = sap_szt.get(idx, 0) + szt
            sap_jm[idx] = jm

            if sap_user is None and "USER" in df_sap.columns:
                sap_user = str(r["USER"]).strip()
            if sap_date is None and "DATA" in df_sap.columns:
                sap_date = str(r["DATA"]).strip()


        # --- budowa raportu w kolejności Hydry ---
        seq_set = set(seq)
        
        # --- WERSJA DOCZELOWA: nie pokazuj powtórek jeśli ilość ta sama ---
        lines = []
        extras_not_in_sap = []
        missing_seen = set() # by nie duplikować wpisów
        
        shown_qty: dict[str, float] = {}   # INDEKS -> ostatnia pokazana ilość
        rows = []
        lp = 1

        for gp in seq:
            if gp not in sap_qty:
                ...
                continue

            qty = float(sap_qty.get(gp, 0.0))
            szt = int(sap_szt.get(gp, 0))
            jm = sap_jm.get(gp, "M")

            if gp in shown_qty and abs(shown_qty[gp] - qty) < 1e-9:
                continue

            shown_qty[gp] = qty
            lines.append(f"{lp:>2}. {gp:<18}  {qty:>10.1f} {jm:<2}  {szt:>6}")

            rows.append({
                "lp": lp,
                "index": gp,
                "qty_m": f"{qty:.1f} {jm}",
                "pcs": f"{szt}",
                "pallets": "",
            })

            lp += 1

        # --- pozycje w SAP, których nie ma w Hydrze (kolejność nieznana) ---
        missing_in_hydra = [idx for idx in sap_qty.keys() if idx not in seq_set]

        header = "RAPORT PODSTAW POD OKLEJANIE (SAP ułożony wg Hydry)\n"
        meta = []
        if sap_date: meta.append(f"Data: {sap_date}")
        if sap_user: meta.append(f"Użytkownik: {sap_user}")
        # linia_value ustawiasz w GUI albo na sztywno
        meta.append(f"Linia: {linia_value}")
        meta_txt = " | ".join(meta) + "\n\n"

        report_text = header + meta_txt
        report_text += "LP  INDEKS               ILOŚĆ       M   SZT\n"
        report_text += "-" * 48 + "\n"
        report_text += "\n".join(lines)

        if extras_not_in_sap:
            report_text += "\n\nW Hydrze, ale BRAK w SAP:\n"
            report_text += "\n".join(f"- {x}" for x in extras_not_in_sap)

        if missing_in_hydra:
            report_text += "\n\nW SAP, ale BRAK w Hydrze (kolejność nieznana):\n"
            report_text += "\n".join(f"- {x}  {sap_qty[x]:.1f} {sap_jm.get(x,'M')}" for x in missing_in_hydra)
            
        text.configure(state="normal")
        text.delete("1.0", "end")
        text.insert("end", report_text)
        text.configure(state="disabled")
        _upadate_placeholder_visibility()
        
        app_state["last_report_text"] = report_text
        app_state["last_report_kind"] = "sap"

        # Dane strukturalne pod DOCX (layout jak w Wordzie)
        end_by_machine = app_state.get("end_by_machine", {}) or {}
        shift_line = end_by_machine.get(linia_value, "")
        app_state["last_report_data"] = {
            "shift_info": shift_line if shift_line else f"{pl_weekday_name(date.today())} (zmiana 1) ({date.today().strftime('%d.%m.%Y')})",
            "report_date": str(date.today()),
            "user": sap_user or "",
            "line": linia_value,
            "machine": "Maszyna 2",
            "rows": rows
        }

        _set_print_visible(True)   # jeśli masz przycisk druku ukrywany/pokazywany
        

     # --- funkcje eksportu/drukowania/edycji raportu DOCX ---       
    def export_report_docx(report_data: dict, template_path: Path | None = None) -> Path:
        """Generuje DOCX na bazie template. Zwraca ścieżkę do wygenerowanego pliku."""
        if not report_data:
            raise ValueError("Brak danych raportu (report_data).")

        if template_path is None:
            template_path = BASE_DIR / "templates" / "report_template.docx"

        if not template_path.exists():
            raise FileNotFoundError(f"Brak szablonu DOCX: {template_path}")
        
        doc = Document(str(template_path))

        def replace_all(old: str, new: str) -> None:
            # paragrafy
            for p in doc.paragraphs:
                if old in p.text:
                    for r in p.runs:
                        r.text = r.text.replace(old, new)
            # tabele
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if old in cell.text:
                            cell.text = cell.text.replace(old, new)

        # 1) Podmień pola nagłówka
        replace_all("{{SHIFT_INFO}}", str(report_data.get("shift_info", "")))
        replace_all("{{REPORT_DATE}}", str(report_data.get("report_date", "")))
        replace_all("{{USER}}", str(report_data.get("user", "")))
        replace_all("{{LINE}}", str(report_data.get("line", "")))
        replace_all("{{MACHINE}}", str(report_data.get("machine", "")))
        replace_all("{{PALLETS_TOTAL}}", str(report_data.get("pallets_total", "")))

        # 2) Wypełnij tabelę: znajdź wiersz-placeholder i zastąp go danymi
        rows = report_data.get("rows", []) or []
        if not rows:
            # nic do tabeli, zostaw placeholdery jako puste
            replace_all("{{ROW_LP}}", "")
            replace_all("{{ROW_INDEX}}", "")
            replace_all("{{ROW_QTY_M}}", "")
            replace_all("{{ROW_PCS}}", "")
            replace_all("{{ROW_PALLETS}}", "")
        else:
            # szukamy pierwszej tabeli z placeholderem
            target_table = None
            placeholder_row_idx = None
            for t in doc.tables:
                for i, r in enumerate(t.rows):
                    if any("{{ROW_LP}}" in c.text for c in r.cells):
                        target_table = t
                        placeholder_row_idx = i
                        break
                if target_table is not None:
                    break

            if target_table is None or placeholder_row_idx is None:
                raise ValueError("Nie znalazłem w szablonie wiersza placeholderów ({{ROW_LP}}...).")

            # usuń wiersz placeholder
            tbl = target_table._tbl
            tr = target_table.rows[placeholder_row_idx]._tr
            tbl.remove(tr)

            # dodaj wiersze danych
            for item in rows:
                r = target_table.add_row().cells
                r[0].text = str(item.get("lp", ""))
                r[1].text = str(item.get("index", ""))
                r[2].text = str(item.get("qty_m", ""))
                r[3].text = str(item.get("pcs", ""))
                r[4].text = str(item.get("pallets", ""))

        # 3) Zapis
        out_dir = Path(tempfile.gettempdir()) / "production_counter_reports"
        out_dir.mkdir(parents=True, exist_ok=True)
        stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
        safe_line = str(report_data.get("line", "LINIA")).replace("/", "-")
        out_path = out_dir / f"Zapotrzebowanie_{safe_line}_{stamp}.docx"
        doc.save(str(out_path))
        return out_path

    # helper: otwiera plik DOCX w domyślnej aplikacji
    def open_docx(path: Path) -> None:
        try:
            os.startfile(str(path))  # Windows
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się otworzyć pliku:\n{e}")


    # helper: drukuje plik DOCX
    def print_docx(path: Path) -> None:
        try:
            os.startfile(str(path), "print")  # Windows
        except Exception as e:
            messagebox.showerror("Błąd druku", f"Nie udało się uruchomić druku:{e}")


    # helper: drukuje aktualny raport
    def print_current_report() -> None:
        kind = app_state.get("last_report_kind")  # "db" albo "sap"
        report_text_full = app_state.get("last_report_text", "")

        # 1) SAP -> druk DOCX (jeśli masz last_report_data)
        if kind == "sap" and app_state.get("last_report_data"):
            try:
                docx_path = export_report_docx(app_state["last_report_data"])
                print_docx(docx_path)
            except Exception as e:
                messagebox.showerror("Błąd druku", f"Nie udało się drukować DOCX:\n{e}")
            return

        # 2) DB -> drukuj skrót (jak w starej wersji)
        if kind == "db":
            report_to_print = make_print_summary(report_text_full)
        else:
            # fallback: drukuj to co jest (np. inne tryby)
            report_to_print = report_text_full

        if not report_to_print.strip():
            messagebox.showwarning("Brak raportu", "Nie ma nic do wydrukowania.")
            _set_print_visible(False)
            return

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8") as f:
                f.write(report_to_print)
                path = f.name
            os.startfile(path, "print")  # Windows
        except Exception as e:
            messagebox.showerror("Błąd druku", f"Nie udało się uruchomić druku:\n{e}")

    # helper: edytuje aktualny raport (tylko SAP)
    def edit_current_report() -> None:
        report_data = app_state.get("last_report_data")
        if not report_data:
            messagebox.showwarning("Brak raportu", "Najpierw wygeneruj raport.")
            return

        try:
            docx_path = export_report_docx(report_data)
            open_docx(docx_path)
        except Exception as e:
            messagebox.showerror("Błąd edycji", f"Nie udało się przygotować/otworzyć DOCX:\n{e}")
      

    # placeholder (nakładana etykieta wewnątrz Textboxa)
    placeholder_text = "Program do obliczania produkcji \n\n1. Kliknij „Wczytaj maszyny”, aby pobrać aktualne zlecenia z bazy danych.\n\n 2. Wybierz maszyny do przeliczenia i ustaw sztuki na zmianę. \n\n 3. Kliknij „Przelicz produkcję”, aby sprawdzić, \n\n do której zmiany i dnia potrwa produkcja. \n\n Opcjonalnie: \n\n – „Wczytaj plik” umożliwia przeliczenie produkcji z pliku Excel (.xlsx, .xls), \n\n jeśli baza danych jest niedostępna."
    placeholder_lbl = customtkinter.CTkLabel(text, text=placeholder_text, justify="center", text_color="#888888", font=default_font)
    # umieść placeholder wewnątrz textboxa, wyśrodkowany
    placeholder_lbl.place(in_=text, relx=0.5, rely=0.5, anchor="center")

    # zmienna statusu i etykieta (wyświetlają liczbę wczytanych rekordów lub komunikaty)
    result_var = tk.StringVar(value="")
    status_label = customtkinter.CTkLabel(
        left,
        textvariable=result_var,
        anchor="w",
        justify="left",
        wraplength=220  # dopasuj do szerokości lewego panelu
    )
    status_label.grid(row=97, column=0, padx=6, pady=(6, 6), sticky="ew")  

    # helper: tworzy skrócony raport do druku z pełnego raportu DB
    def make_print_summary(report_text: str) -> str:
        """
        Z pełnego raportu wycina tylko:
        === WLO-... ===
        Produkcja będzie trwać do: ...
        """
        if not report_text or not report_text.strip():
            return ""

        lines = report_text.splitlines()
        out: list[str] = []
        current_machine: str | None = None
        current_end: str | None = None

        for line in lines:
            line = line.strip()

            # nagłówek maszyny
            m = re.match(r"^===\s*(WLO-[A-Z]\d{3})\s*===$", line)
            if m:
                # jeśli kończymy poprzednią maszynę
                if current_machine:
                    out.append(f"=== {current_machine} ===")
                    out.append(current_end or "Produkcja będzie trwać do: brak danych")
                    out.append("")  # pusta linia między maszynami
                current_machine = m.group(1)
                current_end = None
                continue

            # linia końca produkcji
            if line.startswith("Produkcja będzie trwać do:"):
                current_end = line

            # jeśli maszyna nie ma danych
            if line == "Brak danych." and current_machine:
                current_end = "Produkcja będzie trwać do: brak danych"

        # domknij ostatnią maszynę
        if current_machine:
            out.append(f"=== {current_machine} ===")
            out.append(current_end or "Produkcja będzie trwać do: brak danych")
            out.append("")

        title = "---- Przewidywane zakończenie produkcji --- \n\n"
        return title + "\n".join(out).rstrip() + "\n"    


    # helper: aktualizuje widoczność placeholdera w textboxie
    def _upadate_placeholder_visibility():
        # sprawdź zawartość i pokaż/ukryj placeholder
        content = text.get("1.0", "end-1c")
        if content.strip():
            placeholder_lbl.place_forget()
        else:
            placeholder_lbl.place(in_=text, relx=0.5, rely=0.5, anchor="center")
            
    # funkcja pokazująca popup wyboru maszyn
    def show_machine_select_popup(machines: list[str], on_confirm):
        popup = ctk.CTkToplevel(root)
        popup.title("Wybór maszyn")
        popup.geometry("620x460")
        popup.resizable(False, False)
        center_popup(root, popup)
        popup.grab_set()
        
        df_mc: pd.DataFrame | None = None  # <-- KLUCZOWA LINIA
        
        # wczytaj machine_config (z cache app_state jeśli jest)
        try:
            df_mc = app_state.get("machine_cfg")
            if df_mc is None or df_mc.empty:
                df_mc = load_machine_config()
                app_state["machine_cfg"] = df_mc
        except Exception as e:
            messagebox.showerror("Błąd konfiguracji", f"Nie mogę wczytać machine_config.csv:\n{e}")
            popup.destroy()
            return

        # tu już NA PEWNO jest DataFrame
        if df_mc is None:
            messagebox.showerror("Błąd", "Brak konfiguracji maszyn (df_mc=None).")
            popup.destroy()
            return

        df_mc_df: pd.DataFrame = df_mc

        vars_map: dict[str, tk.BooleanVar] = {}
        pps_vars: dict[str, tk.StringVar] = {}  # pieces per shift (entry)
        default_pps: dict[str, int] = {}

        title = ctk.CTkLabel(
            popup,
            text="Wybierz maszyny do przeliczenia",
            font=ctk.CTkFont(size=16, weight="bold"),
        )
        title.pack(pady=(12, 8))

        header = ctk.CTkFrame(popup, fg_color="transparent")
        header.pack(fill="x", padx=12)

        ctk.CTkLabel(header, text="Maszyna", width=220, anchor="w").pack(side="left")
        ctk.CTkLabel(header, text="Szt./zmianę", width=120, anchor="w").pack(side="left")

        scroll = ctk.CTkScrollableFrame(popup, width=580, height=280)
        scroll.pack(padx=12, pady=8, fill="both", expand=True)

        # helper: weź default szt./zmianę z machine_config.csv
        def get_pps_for(machine: str) -> int:
            row = df_mc_df.loc[df_mc_df["workplace"].astype("string").str.strip() == str(machine).strip()]
            if row.empty:
                return 0
            return int(row.iloc[0]["count_by_shift"])

        # budujemy listę: checkbox + entry
        for m in machines:
            row = ctk.CTkFrame(scroll, fg_color="transparent")
            row.pack(fill="x", padx=6, pady=3)

            var = tk.BooleanVar(value=False)
            vars_map[m] = var

            cb = ctk.CTkCheckBox(row, text=m, variable=var, width=220)
            cb.pack(side="left", padx=(4, 10))

            pps = get_pps_for(m)
            default_pps[m] = pps
            sv = tk.StringVar(value=str(pps))
            pps_vars[m] = sv

            entry = ctk.CTkEntry(row, width=120, textvariable=sv)
            entry.pack(side="left")

            ctk.CTkLabel(row, text="szt./zmianę", text_color="#aaaaaa").pack(side="left", padx=8)

        # --- TOGGLE: Wybierz/Odznacz wszystkie ---
        def _all_selected() -> bool:
            return all(v.get() for v in vars_map.values()) if vars_map else False

        # helper: odśwież napis przycisku
        def _refresh_toggle_btn_text():
            toggle_btn.configure(text="Odznacz wszystkie" if _all_selected() else "Wybierz wszystkie")

        # funkcja toggle wybiera/odznacza wszystkie
        def toggle_select_all():
            if _all_selected():
                for v in vars_map.values():
                    v.set(False)
            else:
                for v in vars_map.values():
                    v.set(True)
            _refresh_toggle_btn_text()

        # gdy user klika pojedyncze checkboxy, aktualizuj napis przycisku
        for v in vars_map.values():
            try:
                v.trace_add("write", lambda *args: _refresh_toggle_btn_text())
            except Exception:
                pass
         
         # helper: parsowanie int z tekstu (lub None)       
        def _parse_int_or_none(s: str):
            s = (s or "").strip().replace(",", ".")
            if s == "":
                return None
            # pozwól wpisać "120" albo "120.0"
            try:
                return int(float(s))
            except Exception:
                return None

        # potwierdzenie wyboru
        def confirm():
            nonlocal df_mc_df
            selected = [m for m, v in vars_map.items() if v.get()]
            if not selected:
                messagebox.showwarning("Brak wyboru", "Zaznacz przynajmniej jedną maszynę.")
                return

            pps_by_machine = {}
            for m in selected:
                val = _parse_int_or_none(pps_vars[m].get())
                if val is None or val <= 0:
                    messagebox.showerror("Błędna wartość", f"Maszyna {m}: szt./zmianę musi być > 0.")
                    return
                pps_by_machine[m] = int(val)

            # 1) sprawdź zmiany szt./zmianę i ewentualnie zapisz do machine_config.csv
            changes = []
            for m, sv in pps_vars.items():
                new_val = _parse_int_or_none(sv.get())
                if new_val is None or new_val < 0:
                    messagebox.showerror("Błędna wartość", f"Maszyna {m}: nieprawidłowa wartość szt./zmianę.")
                    return
                old_val = int(default_pps.get(m, 0))
                if new_val != old_val:
                    changes.append((m, old_val, new_val))

            if changes:
                # pytamy 1 raz zbiorczo (czytelniej niż spam popupami)
                preview = "\n".join([f"{m}: {old} → {new}" for (m, old, new) in changes[:12]])
                if len(changes) > 12:
                    preview += "\n..."

                if messagebox.askyesno(
                    "Zapis do konfiguracji",
                    "Wykryto zmiany szt./zmianę:\n\n"
                    f"{preview}\n\n"
                    "Zapisać do machine_config.csv?"
                ):
                    for m, _, new_val in changes:
                        mask = df_mc_df["workplace"].astype("string").str.strip() == str(m).strip()
                        if mask.any():
                            df_mc_df.loc[mask, "count_by_shift"] = int(new_val)
                        else:
                            # jeśli maszyny nie ma w configu – dopisujemy z zerową prędkością (żeby nie wywaliło)    
                            new_row = pd.DataFrame([{
                                "workplace": str(m).strip(),
                                "speed_m_per_min": 0.0,
                                "count_by_shift": int(new_val),
                            }])

                            df_mc_df = pd.concat([df_mc_df, new_row], ignore_index=True)

                    app_state["machine_cfg"] = df_mc_df
                    save_machine_config(df_mc_df, MACHINE_CONFIG_PATH)
                    
                    # spróbuj zaktualizować DB jeśli DB jest źródłem
                    if app_state.get("machine_cfg_source") in ("db", "db+csv"):
                        try:
                            for m, _, new_val in changes:
                                update_count_by_shift(m, new_val)
                        except Exception as e:
                            messagebox.showwarning(
                                "Uwaga",
                                "Zapisano do pliku, ale nie udało się zaktualizować bazy (brak uprawnień lub błąd połączenia).\n\n"
                                f"Szczegóły: {e}"
                            )
                    df_cfg, source, missing = merge_db_and_csv_config(sync_missing_to_db=False)
                    app_state["machine_cfg"] = df_cfg
                    app_state["machine_cfg_source"] = source
                    
            popup.destroy()
            on_confirm(selected, pps_by_machine)

        btn_frame = ctk.CTkFrame(popup, fg_color="transparent")
        btn_frame.pack(fill="x", padx=12, pady=12)

        toggle_btn = ctk.CTkButton(btn_frame, text="Wybierz wszystkie", command=toggle_select_all)
        toggle_btn.pack(side="left")

        ctk.CTkButton(btn_frame, text="Anuluj", command=popup.destroy).pack(side="right", padx=(8, 0))
        ctk.CTkButton(btn_frame, text="Przelicz produkcję", command=confirm).pack(side="right")

        _refresh_toggle_btn_text()

    # funkcja budująca raport tekstowy z DB        
    def build_db_report_pieces(
        df: pd.DataFrame,
        df_cfg: pd.DataFrame,
        selected_machines: list[str],
        pps_by_machine: dict[str, int],
        start_d: date,
        start_shift: int,
        include_weekends: bool,
    ) -> str:
        lines = []
        lines.append("---- Przewidywane zakończenie produkcji --- \n")

        for machine in selected_machines:
            df_one = df[df["workplace"] == machine].copy()
            if df_one.empty:
                lines.append(f"=== {machine} ===")
                lines.append("Brak danych.\n")
                continue

            # soll i gut w sztukach
            soll = pd.to_numeric(df_one["target_value_pcs"], errors="coerce").fillna(0)
            gut = pd.to_numeric(df_one["good_qty_pcs"], errors="coerce").fillna(0)

            remaining = (soll - gut).clip(lower=0)
            total_remaining = float(remaining.sum())

            pps = int(pps_by_machine.get(machine, 0))

            lines.append(f"=== {machine} ===")
            lines.append(f"Pozostało: {total_remaining:.0f} szt.")

            if pps <= 0:
                lines.append("Szt./zmianę: BRAK / 0 (nie da się policzyć zmian)\n")
                continue

            # --- ZBROJENIA (tak samo jak Excel) ---
            # normalizacja kluczy
            df_one["profile"] = df_one["profile"].astype("string").str.strip()
            df_one["side"] = df_one["side"].astype("string").str.strip().str.zfill(4)

            cfg = df_cfg.copy()
            cfg["profile"] = cfg["profile"].astype("string").str.strip()
            cfg["side"] = cfg["side"].astype("string").str.strip().str.zfill(4)

            # merge setting_time
            df_one = df_one.merge(
                cfg[["profile", "side", "setting_time"]],
                on=["profile", "side"],
                how="left",
            )

            df_one["setting_time"] = pd.to_numeric(df_one["setting_time"], errors="coerce").fillna(0).astype(int)

            # zasada: dla strony 0020 nie liczymy zbrojeń
            df_one.loc[df_one["side"] == "0020", "setting_time"] = 0

            # liczymy unikalne zbrojenia (profile+side)
            unique_setups = df_one[["profile", "side", "setting_time"]].drop_duplicates()
            setup_count = int((unique_setups["setting_time"] > 0).sum())
            setup_min = float(unique_setups["setting_time"].sum())
            setup_shifts = setup_min / (8 * 60)  # 8h = 480 min

            # --- ZMIANY: produkcja + zbrojenia ---
            prod_shifts = total_remaining / pps if total_remaining > 0 else 0.0
            shifts_exact = prod_shifts + setup_shifts
            shifts_rounded = round_shifts_custom(shifts_exact)

            # koniec produkcji dla tej maszyny
            end_d, end_s = add_shifts(
                start_date=start_d,
                start_shift=start_shift,
                shifts_count=shifts_rounded,
                include_weekends=include_weekends,
            )

            lines.append(f"Szt./zmianę: {pps}")
            lines.append(f"Ilość zbrojeń profili: {setup_count}")
            lines.append(f"Czas zbrojeń: {setup_min:.0f} min")
            lines.append(f"Zmiany (8h): {shifts_exact:.2f} → {shifts_rounded}")
            lines.append(f"Start liczenia: {pl_weekday_name(start_d)} ({start_d.isoformat()}) zmiana {start_shift}")
            lines.append("---------------------------------------------------------------------")
            lines.append(f"Produkcja będzie trwać do: {pl_weekday_name(end_d)} (zmiana {end_s}) ({end_d.strftime('%d.%m.%Y')})\n")

        return "\n".join(lines)
       
    # funkcja przeliczająca produkcję z DB i pokazująca wyniki w textboxie    
    def calculate_from_db(selected_machines, pps_by_machine):
        # 1) parametry czasu (ten sam popup co w Excelu)
        # bierzemy default szt./zmianę z pierwszej zaznaczonej maszyny
        default_pps = int(pps_by_machine.get(selected_machines[0], 0)) if selected_machines else 0

        choice = ask_schedule_popup(root) 
        
        if not choice:
            return  # anulowano

        calendar_mode = choice.get("calendar", "workdays")
        include_weekends = (calendar_mode == "all")
        start_shift = int(choice.get("start_shift", 1))

        start_mode = choice.get("start_mode", "today")
        start_date_str = choice.get("start_date", date.today().isoformat())

        if start_mode == "date":
            try:
                start_d = datetime.strptime(start_date_str, "%Y-%m-%d").date()
            except ValueError:
                messagebox.showerror("Błąd daty", "Podaj datę w formacie YYYY-MM-DD.")
                return
        else:
            start_d = date.today()

        # 2) dane z DB
        df_raw = fetch_orders_for_machines(selected_machines)
        df = normalize_db_df(df_raw)
        
        # 2b) wczytaj profile_config (z cache app_state jeśli jest)
        try:
            df_cfg = app_state.get("cfg")
            if df_cfg is None or df_cfg.empty:
                df_cfg = load_profile_confing()
                app_state["cfg"] = df_cfg
        except Exception as e:
            messagebox.showerror("Błąd konfiguracji", f"Nie mogę wczytać profile_config.csv:\n{e}")
            return

        # 3) raport
        report = build_db_report_pieces(
            df=df,
            df_cfg=df_cfg,
            selected_machines=selected_machines,
            pps_by_machine=pps_by_machine,
            start_d=start_d,
            start_shift=start_shift,
            include_weekends=include_weekends,
        )
        show_text_view()  # <- bez argumentów

        text.configure(state="normal")
        text.delete("1.0", "end")
        text.insert("end", report)
        text.configure(state="disabled")
        _upadate_placeholder_visibility()

        app_state["last_report_text"] = report
        app_state["last_report_kind"] = "db"
        app_state["last_report_data"] = None  # żeby nie próbował DOCX
        _set_print_visible(True)
        
        # --- wyciągnij z raportu DB daty zakończenia dla maszyn ---
        end_by_machine = {}
        current_machine = None

        for line in report.splitlines():
            m = re.match(r"^===\s*(.+?)\s*===$", line.strip())
            if m:
                current_machine = m.group(1).strip()
                continue

            if current_machine and line.strip().startswith("Produkcja będzie trwać do:"):
                end_by_machine[current_machine] = line.strip()

        app_state["end_by_machine"] = end_by_machine


    # funkcja wczytująca dane maszyn z DB i pokazująca popup wyboru maszyn
    def loading_machine_data(parent):
        # wczytaj konfigurację maszyn (DB + CSV)
        df_cfg, source, missing = merge_db_and_csv_config(sync_missing_to_db=False)
        app_state["machine_cfg"] = df_cfg
        app_state["machine_cfg_source"] = source
                
        try:
            machines = fetch_available_machines()
            if source == "db+csv":
                # lepiej status w pasku, ale możesz też messagebox (tylko raz)
                messagebox.showerror(f"Konfiguracja: DB + CSV (brak w DB: {len(missing)})")                     
        except Exception as e:
            messagebox.showerror("Brak sterownika ODBC / brak dostępu do sieci firmowej \n\n Możesz użyć trybu: Wczytaj plik (Excel", str(e))
            return

        show_machine_select_popup(machines, calculate_from_db)
  
    # funkcja przełączania motywu
    def change():
        if customtkinter.get_appearance_mode() == "Light":
            customtkinter.set_appearance_mode("Dark")
            try:
                ch_theme.configure(text="Jasny motyw")
            except Exception:
                pass
        else:
            customtkinter.set_appearance_mode("Light")
            try:
                ch_theme.configure(text="Ciemny motyw")
            except Exception:
                pass
    
    # funkcja wyświetlania tylko tekstu w textboxie (ukrywa tabelę, jeśli była)
    def show_text_view():
        tf = app_state.get("table_frame")
        if tf is not None:
            try:
                tf.destroy()
            except Exception:
                pass
        app_state["table_frame"] = None
        
        
    # funkcja wyświetlania DataFrame w tabeli Treeview
    def show_table_from_df(df: pd.DataFrame):
        # 1) schowaj textbox
        # usuń poprzednią, jeśli istnieje
        old_tf = app_state.get("table_frame")
        if old_tf is not None:
            try:
                old_tf.destroy()
            except Exception:
                pass

        # 3) zbuduj nową ramkę + tree
        tf = customtkinter.CTkFrame(right)
        tf.grid(row=1, column=0, sticky="nsew")
        app_state["table_frame"] = tf

        cols = list(map(str, df.columns))
        tree = ttk.Treeview(tf, columns=cols, show="headings")

        vsb = ttk.Scrollbar(tf, orient="vertical", command=tree.yview)
        hsb = ttk.Scrollbar(tf, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        for col in cols:
            tree.heading(col, text=col)
            tree.column(col, width=120, anchor="w")

        # Uwaga: nie ładuj bez limitu tysięcy wierszy – daj podgląd
        for i, row in enumerate(df.itertuples(index=False, name=None)):
            if i >= 2000:
                break
            tree.insert("", "end", values=row)

        tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        tf.grid_rowconfigure(0, weight=1)
        tf.grid_columnconfigure(0, weight=1)


    # licze dni tygodnia z niestandardowym zaokrągleniem
    def pl_weekday_name(d: date) -> str:
        names = ["poniedziałek", "wtorek", "środa", "czwartek", "piątek", "sobota", "niedziela"]
        return names[d.weekday()]

    def round_shifts_custom(shifts: float) -> int:
        """4.5 -> 4 (w dół), 4.7 -> 5 (w górę)."""
        frac = shifts - math.floor(shifts)
        return math.floor(shifts) if frac < 0.5 else math.ceil(shifts)

    def next_workday(d: date) -> date:
        d += timedelta(days=1)
        while d.weekday() >= 5:  # 5=sob, 6=nd
            d += timedelta(days=1)
        return d
    
    # zwraca liczbę zmian danego dnia
    def shifts_per_day_for_date(d: date, include_weekends: bool) -> int:
        """
        Zwraca liczbę zmian danego dnia:
        - pn-pt: 3 zmiany
        - sob-nd: 1 zmiana (tylko jeśli include_weekends=True)
        - sob-nd przy include_weekends=False i tak nie powinny wystąpić (bo skaczemy po workday)
        """
        if include_weekends and d.weekday() >= 5:  # 5=sob, 6=nd
            return 1
        return SHIFTS_PER_DAY  # u Ciebie = 3

    # zwraca (end_date, end_shift) po wykonaniu shifts_count zmian
    def add_shifts(start_date: date, start_shift: int, shifts_count: int, include_weekends: bool) -> tuple[date, int]:
        """
        Zwraca (end_date, end_shift) po wykonaniu shifts_count zmian,
        startując od start_date i start_shift.

        Zasada: jeśli include_weekends=True to w sob/nd jest tylko 1 zmiana.
        """
        if shifts_count <= 0:
            return start_date, start_shift

        d = start_date

        # jeżeli NIE liczymy weekendów, to start_date w weekend przepchnij na poniedziałek
        if not include_weekends and d.weekday() >= 5:
            d = next_workday(d)

        # normalizacja start_shift (np. start w weekend i ktoś wybrał zmianę 2/3 -> spada do 1)
        max_shifts_today = shifts_per_day_for_date(d, include_weekends)
        s = int(start_shift)
        if s < 1:
            s = 1
        if s > max_shifts_today:
            s = 1

        # koniec jest na shifts_count-tej zmianie, więc przesuwamy slot shifts_count-1 razy
        moves = shifts_count - 1

        for _ in range(moves):
            max_shifts_today = shifts_per_day_for_date(d, include_weekends)

            s += 1
            if s > max_shifts_today:
                # przechodzimy na następny dzień
                s = 1
                if include_weekends:
                    d += timedelta(days=1)
                else:
                    d = next_workday(d)

                # po zmianie dnia, jeśli weekendy wliczamy, max zmian może być 1
                max_shifts_today = shifts_per_day_for_date(d, include_weekends)
                if s > max_shifts_today:
                    s = 1

        return d, s

    # wczytuje konfigurację maszyn
    def load_machine_config() -> pd.DataFrame:
        if not MACHINE_CONFIG_PATH.exists():
            raise FileNotFoundError(f"Plik konfiguracyjny nie istnieje: {MACHINE_CONFIG_PATH}")

        df_mc = pd.read_csv(MACHINE_CONFIG_PATH, sep=";", encoding="utf-8", dtype={"workplace": "string"})
        df_mc.columns = [col.strip().strip('"').rstrip(";") for col in df_mc.columns]
        
        df_mc.columns = df_mc.columns.str.strip()
        if "machine" in df_mc.columns and "workplace" not in df_mc.columns:
            df_mc = df_mc.rename(columns={"machine": "workplace"})


        required = {"workplace", "speed_m_per_min", "count_by_shift"}
        missing = required - set(df_mc.columns)
        if missing:
            raise ValueError(f"Brakujące kolumny w machine_config.csv: {missing}")

        df_mc["workplace"] = df_mc["workplace"].astype("string").str.strip()
        df_mc["speed_m_per_min"] = pd.to_numeric(df_mc["speed_m_per_min"], errors="coerce").fillna(0.0)
        df_mc["count_by_shift"] = pd.to_numeric(df_mc["count_by_shift"], errors="coerce").fillna(0).astype(int)

        # workplace powinno być unikalne
        duplicates = df_mc.duplicated(subset=["workplace"], keep=False)
        if duplicates.any():
            dup_rows = df_mc[duplicates]
            raise ValueError(f"Duplikaty w machine_config.csv (workplace musi być unikalne):\n{dup_rows}")

        return df_mc

    # sprawdz czy plik istnieje i wczytaj konfigurację profili
    def load_profile_confing() -> pd.DataFrame:
        # spradz czy plik istnieje
        if not CONFING_PATH.exists():
            raise FileNotFoundError(f"Plik konfiguracyjny nie istnieje: {CONFING_PATH}")

        # wczytaj plik konfiguracyjny
        df_cfg = pd.read_csv(
            CONFING_PATH,
            sep=";",
            encoding="utf-8",
            dtype={"profile": "string", "side": "string"},
        )
        # na wszelki wypadek: usunięcie spacji z nazw kolumn
        df_cfg.columns = [col.strip().strip('"').rstrip(";") for col in df_cfg.columns]

        required = {"profile", "side", "setting_time"}
        missing = required - set(df_cfg.columns)
        if missing:
            raise ValueError(f"Brakujące kolumny w pliku konfiguracyjnym: {missing}") 

        # upewnij się, że setting_time jest liczbą
        df_cfg["setting_time"] = pd.to_numeric(df_cfg["setting_time"], errors="coerce").fillna(0).astype(int)

        # kontrola duplikatów (profile+side musi być 1:1)
        duplicates = df_cfg.duplicated(subset=["profile", "side"], keep=False)
        if duplicates.any(): 
            dup_rows = df_cfg[duplicates]
            raise ValueError(f"Duplikaty w pliku konfiguracyjnym (profile+side musi być unikalne):\n{dup_rows}")

        return df_cfg
    
    #--- POPUP: wybór trybu przeliczenia produkcji ---
    def ask_calc_mode_popup(parent, workplace: str, default_speed: float, default_pieces_per_shift: int):
        result: Dict[str, Optional[Dict[str, Any]]] = {"value": None}
        
        # dni roboczoe lub weekend /kalendarz
        calendar_var = tk.StringVar(value="workdays")
        start_shift_var = tk.IntVar(value=1) 
        
        #dni tygodnia / kalendarz
        start_mode_var = tk.StringVar(value="today")  # "today" albo "date"
        start_date_var = tk.StringVar(value=date.today().isoformat())
    

        # ustaw wyraźne tło dla popupu (uniezależnienie od motywu)
        win = customtkinter.CTkToplevel(parent, fg_color="#2b2b2b")
        win.title("Parametry przeliczenia produkcji")

        try:
            win.transient(parent)
        except Exception:
            pass
        try:
            win.grab_release()
        except Exception:
            pass
        try:
            win.lift()
            win.attributes("-topmost", True)
        except Exception:
            pass

        customtkinter.CTkLabel(
            win,
            text=f"Stanowisko: {workplace}",
            font=customtkinter.CTkFont(size=16, weight="bold"),
            text_color="#eaeaea",
        ).pack(padx=16, pady=(14, 6), anchor="w")

        mode_var = tk.StringVar(value="shift")  #  aktywny tryb"speed" albo "shift"

        frame = customtkinter.CTkFrame(win)
        frame.pack(fill="both", expand=True, padx=16, pady=10)

        # radio + entry szt./zmianę
        row2 = customtkinter.CTkFrame(frame)
        row2.pack(fill="x", padx=10, pady=6)
                
        # --- wiersze z opcjami ---
        # prędkość
        row1 = customtkinter.CTkFrame(frame)
        row1.pack(fill="x", padx=10, pady=(10, 6))
                
        # odstęp wizualny między sekcjami
        spacer = customtkinter.CTkFrame(frame, height=12, fg_color="transparent")
        spacer.pack(fill="x")

        # radio + entry szt./zmianę
        row_cal = customtkinter.CTkFrame(frame)
        row_cal.pack(fill="x", padx=10, pady=(20, 6))
        
        # radio dni robocze / weekendy
        row_start = customtkinter.CTkFrame(frame)
        row_start.pack(fill="x", padx=10, pady=6)
        
        # radio data startu
        row_startdate = customtkinter.CTkFrame(frame)
        row_startdate.pack(fill="x", padx=10, pady=6)
        
        # --- tryb startu daty ---
        customtkinter.CTkLabel(row_startdate, text="Start liczenia:", text_color="#eaeaea").pack(side="left")

        customtkinter.CTkRadioButton(
            row_startdate, text="od dziś", variable=start_mode_var, value="today", text_color="#eaeaea"
        ).pack(side="left", padx=10)

        customtkinter.CTkRadioButton(
            row_startdate, text="od daty", variable=start_mode_var, value="date", text_color="#eaeaea"
        ).pack(side="left", padx=10)

        start_date_entry = customtkinter.CTkEntry(row_startdate, width=130, textvariable=start_date_var)
        start_date_entry.pack(side="left", padx=10)

        customtkinter.CTkLabel(row_startdate, text="(YYYY-MM-DD)", text_color="#aaaaaa").pack(side="left")        
        
        # --- tryb przeliczenia ---
        customtkinter.CTkRadioButton(
            row2, text="Przelicz produkcję poprzez sztuki na zmianę:",
            variable=mode_var, value="shift", text_color="#eaeaea"
        ).pack(side="left")

        pshift_var = tk.StringVar(value=str(default_pieces_per_shift))
        pshift_entry = customtkinter.CTkEntry(row2, width=120, textvariable=pshift_var)
        pshift_entry.pack(side="left", padx=10)
        customtkinter.CTkLabel(row2, text="szt./zmianę", text_color="#eaeaea").pack(side="left")

        customtkinter.CTkLabel(row_start, text="Start od zmiany:", text_color="#eaeaea").pack(side="left")

        customtkinter.CTkRadioButton(
            row_start, text="1", variable=start_shift_var, value=1, text_color="#eaeaea"
        ).pack(side="left", padx=10)

        customtkinter.CTkRadioButton(
            row_start, text="2", variable=start_shift_var, value=2, text_color="#eaeaea"
        ).pack(side="left", padx=10)

        customtkinter.CTkRadioButton(
            row_start, text="3", variable=start_shift_var, value=3, text_color="#eaeaea"
        ).pack(side="left", padx=10)

        # --- tryb przeliczenia ---
        customtkinter.CTkRadioButton(
            row1, text="Przelicz produkcję poprzez prędkość:",
            variable=mode_var, value="speed", text_color="#eaeaea"
        ).pack(side="left")

        speed_var = tk.StringVar(value=str(default_speed))
        speed_entry = customtkinter.CTkEntry(row1, width=120, textvariable=speed_var)
        speed_entry.pack(side="left", padx=10)
        customtkinter.CTkLabel(row1, text="m/min", text_color="#eaeaea").pack(side="left")


        customtkinter.CTkLabel(row_cal, text="Kalendarz:", text_color="#eaeaea").pack(side="left")

        customtkinter.CTkRadioButton(
            row_cal, text="dni robocze", variable=calendar_var, value="workdays", text_color="#eaeaea"
        ).pack(side="left", padx=10)

        customtkinter.CTkRadioButton(
            row_cal, text="dni robocze + weekendy", variable=calendar_var, value="all", text_color="#eaeaea"
        ).pack(side="left", padx=10)
        
        # helper: parsowanie float z tekstu
        def parse_float(s: str) -> float:
            return float(s.replace(",", ".").strip())

        # --- przyciski OK ---
        def on_ok():
            try:
                if mode_var.get() == "speed":
                    v = parse_float(speed_var.get())
                    if v <= 0:
                        raise ValueError("Prędkość musi być > 0.")
                    result["value"] = {
                        "mode": "speed",
                        "speed_m_per_min": v,
                        "calendar": calendar_var.get(),
                        "start_shift": int(start_shift_var.get()),
                        }
                    result["value"]["start_mode"] = start_mode_var.get()
                    result["value"]["start_date"] = start_date_var.get()
                    result["value"]["start_shift"] = int(start_shift_var.get())
                else:
                    v = int(parse_float(pshift_var.get()))
                    if v <= 0:
                        raise ValueError("Sztuki na zmianę muszą być > 0.")
                    result["value"] = {
                        "mode": "shift",
                        "pieces_per_shift": v,
                        "calendar": calendar_var.get(),
                        "start_shift": int(start_shift_var.get()),
                        }
                    result["value"]["start_mode"] = start_mode_var.get()
                    result["value"]["start_date"] = start_date_var.get()
                    result["value"]["start_shift"] = int(start_shift_var.get())

            except Exception as e:
                messagebox.showerror("Błędna wartość", str(e))
                return
            win.destroy()

        # --- przyciski Anuluj ---
        def on_cancel():
            result["value"] = None
            win.destroy()

        btns = customtkinter.CTkFrame(win, fg_color="#2b2b2b")
        btns.pack(fill="x", padx=16, pady=(0, 14))
        customtkinter.CTkButton(btns, text="Anuluj", command=on_cancel, fg_color="#3a3a3a").pack(side="right")
        customtkinter.CTkButton(btns, text="OK", command=on_ok, fg_color="#3a3a3a").pack(side="right", padx=10)

        # ustaw pozycję na środku rodzica przed blokowaniem okna
        try:
            center_popup(parent, win)
        except Exception:
            pass
        parent.wait_window(win)
        return result["value"]

    # note: fallback `tk` popup removed — use `ask_calc_mode_popup` (CTkToplevel)

    # popup do wyboru dla przycisku "Wczytaj maszyny " dni tygodnia i startu 
    def ask_schedule_popup(parent) -> dict | None:
        popup = ctk.CTkToplevel(parent)
        popup.title("Parametry liczenia (DB)")
        popup.resizable(False, False)
        popup.grab_set()

        result: dict | None = None

        # --- Kalendarz ---
        cal_var = tk.StringVar(value="workdays")  # workdays | all

        cal_frame = ctk.CTkFrame(popup)
        cal_frame.pack(fill="x", padx=12, pady=(12, 6))

        ctk.CTkLabel(cal_frame, text="Kalendarz:", width=120, anchor="w").pack(side="left")
        ctk.CTkRadioButton(cal_frame, text="dni robocze", variable=cal_var, value="workdays").pack(side="left", padx=10)
        ctk.CTkRadioButton(cal_frame, text="dni robocze + weekendy", variable=cal_var, value="all").pack(side="left", padx=10)

        # --- Start od zmiany ---
        shift_var = tk.IntVar(value=1)

        shift_frame = ctk.CTkFrame(popup)
        shift_frame.pack(fill="x", padx=12, pady=6)

        ctk.CTkLabel(shift_frame, text="Start od zmiany:", width=120, anchor="w").pack(side="left")
        ctk.CTkRadioButton(shift_frame, text="1", variable=shift_var, value=1).pack(side="left", padx=18)
        ctk.CTkRadioButton(shift_frame, text="2", variable=shift_var, value=2).pack(side="left", padx=18)
        ctk.CTkRadioButton(shift_frame, text="3", variable=shift_var, value=3).pack(side="left", padx=18)

        # --- Start liczenia ---
        start_mode_var = tk.StringVar(value="today")  # today | date
        start_date_var = tk.StringVar(value=date.today().isoformat())

        start_frame = ctk.CTkFrame(popup)
        start_frame.pack(fill="x", padx=12, pady=6)

        ctk.CTkLabel(start_frame, text="Start liczenia:", width=120, anchor="w").pack(side="left")
        ctk.CTkRadioButton(start_frame, text="od dziś", variable=start_mode_var, value="today").pack(side="left", padx=10)
        ctk.CTkRadioButton(start_frame, text="od daty", variable=start_mode_var, value="date").pack(side="left", padx=10)

        date_entry = ctk.CTkEntry(start_frame, width=140, textvariable=start_date_var)
        date_entry.pack(side="left", padx=10)
        ctk.CTkLabel(start_frame, text="(YYYY-MM-DD)", text_color="#aaaaaa").pack(side="left")

        #--- Przycisk OK ---
        def _on_ok():
            nonlocal result
            mode = start_mode_var.get()
            ds = start_date_var.get().strip()

            if mode == "date":
                try:
                    datetime.strptime(ds, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror("Błąd daty", "Podaj datę w formacie YYYY-MM-DD.")
                    return

            result = {
                "calendar": cal_var.get(),
                "start_shift": int(shift_var.get()),
                "start_mode": mode,
                "start_date": ds,
            }
            popup.destroy()

        btn_frame = ctk.CTkFrame(popup, fg_color="transparent")
        btn_frame.pack(fill="x", padx=12, pady=(10, 12))

        ctk.CTkButton(btn_frame, text="Anuluj", command=popup.destroy).pack(side="right", padx=(8, 0))
        ctk.CTkButton(btn_frame, text="OK", command=_on_ok).pack(side="right")

        # ustaw okno na środku rodzica przed blokowaniem
        try:
            center_popup(parent, popup)
        except Exception:
            pass

        popup.wait_window()
        return result

    # główna funkcja przeliczania produkcji
    def calculate_production():
        df_plan = app_state.get("df")
        if df_plan is None or df_plan.empty:
            messagebox.showwarning("Brak danych", "Najpierw wczytaj dane produkcyjne.")
            return

        # config
        try:
            df_cfg = app_state.get("cfg")
            if df_cfg is None or df_cfg.empty:
                df_cfg = load_profile_confing()  # upewnij się, że nazwa funkcji jest poprawna
                app_state["cfg"] = df_cfg
        except Exception as e:
            messagebox.showerror("Błąd wczytywania konfiguracji", str(e))
            return

        df = df_plan.copy()
        
        # --- machine config (speed) ---    
        try:
            df_mc = app_state.get("machine_cfg")
            if df_mc is None or df_mc.empty:
                df_mc = load_machine_config()
                app_state["machine_cfg"] = df_mc
        except Exception as e:
            messagebox.showerror("Błąd wczytywania machine_config.csv", str(e))
            return

        workplaces = df_plan["workplace"].dropna().astype("string").str.strip().unique()
        if len(workplaces) != 1:
            messagebox.showerror("Wybór maszyny", f"W danych wykryto wiele maszyn: {list(workplaces)}")
            return

        workplace = workplaces[0]
        row = df_mc.loc[df_mc["workplace"] == workplace]
        if row.empty:
            messagebox.showerror("Brak konfiguracji", f"Nie znaleziono workplace='{workplace}' w machine_config.csv")
            return

        default_speed = float(row.iloc[0]["speed_m_per_min"])
        default_pieces_per_shift = int(row.iloc[0]["count_by_shift"])


        # Use tkinter fallback popup for debugging visibility issues
        choice = ask_calc_mode_popup(root, workplace, default_speed, default_pieces_per_shift)

        if choice is None:
            return  # Anuluj           

        mask = df_mc["workplace"].astype("string").str.strip() == str(workplace).strip()

        if choice["mode"] == "speed":
            new_speed = float(choice["speed_m_per_min"])

            if abs(new_speed - default_speed) > 1e-9:
                if messagebox.askyesno(
                    "Zapis do konfiguracji",
                    f"Zmieniono prędkość dla {workplace}\n"
                    f"Było: {default_speed}\n"
                    f"Jest: {new_speed}\n\n"
                    "Zapisać do machine_config.csv?"
                ):
                    df_mc.loc[mask, "speed_m_per_min"] = new_speed
                    save_machine_config(df_mc, MACHINE_CONFIG_PATH)
                    app_state["machine_cfg"] = df_mc
                    default_speed = new_speed

        elif choice["mode"] == "shift":  # tryb szt./zmianę
            new_pps = int(choice["pieces_per_shift"])

            if new_pps != default_pieces_per_shift:
                if messagebox.askyesno(
                    "Zapis do konfiguracji",
                    f"Zmieniono szt./zmianę dla {workplace}\n"
                    f"Było: {default_pieces_per_shift}\n"
                    f"Jest: {new_pps}\n\n"
                    "Zapisać do machine_config.csv?"
                ):
                    df_mc.loc[mask, "count_by_shift"] = new_pps
                    save_machine_config(df_mc, MACHINE_CONFIG_PATH)
                    app_state["machine_cfg"] = df_mc
                    default_pieces_per_shift = new_pps

        # --- NORMALIZACJA DANYCH ---
        # wymagane kolumny po normalizacji
        required = {"profile", "side"}
        missing_cols = [c for c in required if c not in df.columns]
        if missing_cols:
            messagebox.showerror("Błąd danych", f"Brak kolumn: {missing_cols}\nSprawdź on_open_file().")
            return

        # normalizacja typów (na wszelki wypadek)
        df["profile"] = df["profile"].astype("string").str.strip()
        df["side"] = df["side"].astype("string").str.strip().str.replace(r"\.0$", "", regex=True).str.zfill(4)

        # upewnij się, że config też ma poprawne typy i kolumny
        for col in ("profile", "side", "setting_time"):
            if col not in df_cfg.columns:
                messagebox.showerror("Błąd configu", f"Brak kolumny '{col}' w profile_config.csv")
                return

        df_cfg = df_cfg.copy()
        df_cfg["profile"] = df_cfg["profile"].astype("string").str.strip()
        df_cfg["side"] = df_cfg["side"].astype("string").str.strip().str.zfill(4)
        df_cfg["setting_time"] = pd.to_numeric(df_cfg["setting_time"], errors="coerce")

        # merge
        df = df.merge(df_cfg[["profile", "side", "setting_time"]], on=["profile", "side"], how="left")

        missing = df[df["setting_time"].isna()]
        if not missing.empty:
            missing_pairs = missing[["profile", "side"]].drop_duplicates().head(20)
            messagebox.showerror(
                "Brak w configu",
                "Nie znaleziono setting_time dla:\n"
                + missing_pairs.to_string(index=False)
                + ("\n..." if len(missing_pairs) >= 20 else "")
            )
            return

        # Twoja reguła biznesowa (jeśli 0020 = brak zbrojenia)
        df.loc[df["side"] == "0020", "setting_time"] = 0

        # --- METRY (suma długości do biegu) ---
        if "target_value_p" not in df.columns or "unit_p" not in df.columns:
            messagebox.showerror("Błąd danych", "Brak kolumn target_value_p lub unit_p – nie mogę policzyć metrów.")
            return

        # --- METRY (pozostałe do wykonania) ---
        target_p = pd.to_numeric(df["target_value_p"], errors="coerce").fillna(0.0)
        good_p = pd.to_numeric(df.get("good_qty_p", 0), errors="coerce").fillna(0.0)

        unit = df["unit_p"].astype("string").str.strip().str.upper()

        # domyślnie: cały target
        remaining_p = target_p.copy()

        # jeśli zlecenie rozpoczęte → odejmij wykonaną część
        mask_started = good_p > 0
        remaining_p.loc[mask_started] = (target_p - good_p).clip(lower=0.0)

        # tylko metry (P w metrach)
        df["length_m"] = remaining_p.where(unit == "M", 0.0)

        total_m = float(df["length_m"].sum())

        # --- SZTUKI (suma sztuk do wykonania) ---
        pieces = pd.to_numeric(df["target_value_s"], errors="coerce").fillna(0.0)
        total_pieces = float(pieces.sum())


        # zbrojenia: unikalna konfiguracja profile+side
        unique_setups = df[["profile", "side", "setting_time"]].drop_duplicates(subset=["profile", "side"])
        real_setups = unique_setups[unique_setups["setting_time"] > 0]
        total_setting_min = float(real_setups["setting_time"].sum())

        # --- CZAS BIEGU (wg trybu z popupu) ---
        total_run_min = 0.0
        run_mode_line = ""

        if choice["mode"] == "speed":
            speed = float(choice["speed_m_per_min"])
            if speed <= 0:
                messagebox.showwarning("Błąd", "Prędkość musi być > 0.")
                return
            total_run_min = total_m / speed
            run_mode_line = f"Tryb biegu: {speed:.2f} m/min\n"
        else:
            pieces_per_shift = int(choice["pieces_per_shift"])
            if pieces_per_shift <= 0:
                messagebox.showwarning("Błąd", "Sztuki na zmianę muszą być > 0.")
                return
            if total_pieces <= 0:
                messagebox.showwarning("Błąd danych", "Suma sztuk (Docelowa wartość (S)) wynosi 0 – nie mogę policzyć trybu shift.")
                return

            shifts_needed = total_pieces / pieces_per_shift
            total_run_min = shifts_needed * 8.0 * 60.0
            run_mode_line = f"Tryb biegu: {pieces_per_shift} szt./zmianę\n"

        # podsumowanie
        total_min = total_setting_min + total_run_min
        total_h = total_min / 60.0
        shifts = total_h / 8.0
        
        calendar_mode = choice.get("calendar", "workdays")
        include_weekends = (calendar_mode == "all")
        start_shift = int(choice.get("start_shift", 1))

        rounded_shifts = round_shifts_custom(shifts)

        start_shift = int(choice.get("start_shift", 1))
        start_mode = choice.get("start_mode", "today")
        start_date_str = choice.get("start_date", date.today().isoformat())

        if start_mode == "date":
            try:
                start_d = datetime.strptime(start_date_str, "%Y-%m-%d").date()
            except ValueError:
                messagebox.showerror("Błąd daty", "Podaj datę w formacie YYYY-MM-DD.")
                return
        else:
            start_d = date.today()

        end_d, end_s = add_shifts(
            start_date=start_d,
            start_shift=start_shift,
            shifts_count=rounded_shifts,
            include_weekends=include_weekends
)
        
        end_line = f"Produkcja będzie trwać do: {pl_weekday_name(end_d)} (zmiana {end_s}) ({end_d.strftime('%d.%m.%Y')})\n"

        result_text = (
            f"Stanowisko:   {workplace}\n"
            f"Pozycje w planie: {len(df)}\n"
            f"Ilość zbrojeń profili: {len(real_setups)}\n"
            f"Suma metrów:  {total_m:.1f} m\n"
            f"Suma sztuk:   {total_pieces:.0f} szt.\n"
            f"{run_mode_line}"
            f"Czas zbrojeń: {total_setting_min:.1f} min\n"
            f"Czas biegu:   {total_run_min:.1f} min\n"
            f"Razem:        {total_min:.1f} min = {total_h:.2f} h\n"
            f"--------------------------------\n"
            f"Zmiany (8h):  {shifts:.2f}\n"
            f"Start liczenia: {pl_weekday_name(start_d)} ({start_d.isoformat()}) zmiana {start_shift}\n"
            f"---------------------------------------------------------------------\n"
            f"{end_line}"
        )

        
        configs = real_setups[["profile", "side", "setting_time"]].sort_values(["profile", "side"])

        result_text += "\nKonfiguracja czasów dla zbrojeń:\n" + configs.to_string(index=False)
        
        text.configure(state="normal")
        text.delete("1.0", "end")
        text.insert("end", result_text)
        text.configure(state="disabled")
        _upadate_placeholder_visibility()
        
        show_text_view()

        # Debug / opcjonalnie podgląd:
        # show_table_from_df(df)
        
     # funkcja zapisu konfiguracji maszyn   
    def save_machine_config(df_mc: pd.DataFrame, path: Path) -> None:
        df_mc.to_csv(path, sep=";", index=False, encoding="utf-8")

    # funkcja wykrywania kolumny strony
    def detect_side_column(df: pd.DataFrame) -> str:
        """
        Zwraca kolumnę, która zawiera strony 20/21/22/23.
        Jeśli nie znajdzie – rzuca wyjątek (dane wejściowe są błędne).
        """
        allowed = {"0020", "0021", "0022", "0023"}

        for col in df.columns:
            s = (
                df[col]
                .astype("string")
                .str.strip()
                .str.replace(r"\.0$", "", regex=True)
                .str.zfill(4)
            )

            non_empty = s[s != ""]
            if non_empty.empty:
                continue

            if non_empty.isin(allowed).mean() > 0.8:
                return col

        raise ValueError("Nie znaleziono kolumny strony (20/21/22/23).")
    # budujesz raport i pokazujesz / zapisujesz



    # funkcja obsługi wczytywania pliku
    def on_open_file():
        file_path = filedialog.askopenfilename(
            title="Wybierz plik",
            filetypes=[
                ("Excel files", ("*.xlsx", "*.xls")),
                ("CSV files", ("*.csv",)),
                ("All files", ("*.*",)),
            ],
        )
        if not file_path:
            return 

        try:
            ext = Path(file_path).suffix.lower()

            if ext in (".xlsx", ".xls"):
                df_raw = pd.read_excel(file_path, engine="openpyxl", header=1)
            elif ext == ".csv":
                df_raw = pd.read_csv(file_path, encoding="utf-8-sig", sep=",", low_memory=False)
            else:
                messagebox.showerror("Błąd", "Nieobsługiwany format pliku. Wybierz .xlsx, .xls lub .csv")
                return

            # 1) czyść nagłówki (strip + normalizacja spacji)
            df_raw.columns = [
                " ".join(str(c).replace("\xa0", " ").strip().split())
                for c in df_raw.columns
            ]
            
            # funkcja pomocnicza do znajdowania kolumny po możliwych nazwach
            def find_col(df_cols, *candidates):
                cols_norm = {str(c).strip(): str(c) for c in df_cols}
                for cand in candidates:
                    # dokładne dopasowanie po strip
                    for k, original in cols_norm.items():
                        if k == cand.strip():
                            return original
                return None
                        
            df_raw.columns = [str(c).strip() for c in df_raw.columns]
            
            good_p_col = find_col(
                df_raw.columns,
                "Ilość dobrej produkcji (P)",
                "Ilość dobrej produkcji(P)",
                "Ilość dobrej produkcji P",
            )
            
            # mapowanie kolumn na podstawie możliwych nazw
            # 2) wybierz stałe kolumny
            needed_fixed = [
                "Stanowisko robocze",
                "Artykuł",
                "Docelowa wartość (P)",
                "Jednostka (P)",
                "Docelowa wartość (S)",
                "Jednostka (S)",
                "Rodzaj zlecenia",
            ]

            # kolumna wykonania jest opcjonalna – jeśli ją znajdziemy, to ją dokładamy
            if good_p_col:
                needed_fixed.insert(3, good_p_col)

            zlecenie_cols = [c for c in df_raw.columns if c.startswith("Zlecenie")]
            if len(zlecenie_cols) < 2:
                raise ValueError("Brakuje drugiej kolumny 'Zlecenie' (tej ze stroną).")

            df = df_raw[needed_fixed + zlecenie_cols].copy()

            # wybierz tę kolumnę Zlecenie*, która jest stroną
            side_col = detect_side_column(df)

            # teraz zostaw tylko potrzebne kolumny + side_col
            df = df[needed_fixed + [side_col]].copy()
            
            # 3) rename kolumn
            rename_map = {
                "Stanowisko robocze": "workplace",
                "Artykuł": "profile",
                "Docelowa wartość (P)": "target_value_p",
                "Jednostka (P)": "unit_p",
                "Docelowa wartość (S)": "target_value_s",
                "Jednostka (S)": "unit_s",
                "Rodzaj zlecenia": "order_type",
                side_col: "side",
            }
            if good_p_col:
                rename_map[good_p_col] = "good_qty_p"

            df = df.rename(columns=rename_map)
            df["good_qty_p"] = pd.to_numeric(df["good_qty_p"], errors="coerce").fillna(0.0)
            
            
            if "good_qty_p" in df.columns:
                df["good_qty_p"] = (
                    df["good_qty_p"]
                    .astype(str)
                    .str.replace("\xa0", "", regex=False)  # twarda spacja
                    .str.replace(" ", "", regex=False)
                    .str.replace(",", ".", regex=False)
                )
                df["good_qty_p"] = pd.to_numeric(df["good_qty_p"], errors="coerce").fillna(0.0)
            else:
                df["good_qty_p"] = 0.0


            # fallback gdy jednak nie było kolumny wykonania
            if "good_qty_p" not in df.columns:
                df["good_qty_p"] = 0.0

            # 4) czyść dane w kolumnach profile i side
            df["profile"] = (
                df["profile"]
                .astype("string")
                .str.strip()
                .str.split("-", n=1)
                .str[0]
            )

            app_state["df"] = df  # zapamiętaj

        except Exception as e:
            messagebox.showerror("Błąd wczytywania pliku", str(e))
            return

        # popup + pokazanie w Treeview (jak u Ciebie)
        popup = customtkinter.CTkToplevel(root)
        popup.title("Wczytano dane")
        popup.transient(root)
        popup.grab_set()

        root.update_idletasks()
        popup.update_idletasks()
        rw = root.winfo_width()
        rh = root.winfo_height()
        rx = root.winfo_rootx()
        ry = root.winfo_rooty()
        pw = popup.winfo_width()
        ph = popup.winfo_height()
        x = rx + (rw - pw) // 2
        y = ry + (rh - ph) // 2
        popup.geometry(f"+{x}+{y}")

        msgbox = f"Wczytano {len(df)} rekordów z pliku:\n{Path(file_path).name}"
        label = customtkinter.CTkLabel(popup, text=msgbox, wraplength=440, anchor="center", justify="center")
        label.pack(padx=12, pady=(12, 6), fill="both")

        confirm_button = customtkinter.CTkButton(
            popup,
            text="OK",
            command=lambda: (show_table_from_df(df), popup.destroy())
        )
        confirm_button.pack(pady=(6, 12))

        root.wait_window(popup)

        # tekstowy preview
        text.configure(state="normal")
        text.delete("1.0", "end")
        text.insert("end", df.head(50).to_string(index=False))
        text.configure(state="disabled")
        _upadate_placeholder_visibility()

    def show_about_popup(parent):
        popup = customtkinter.CTkToplevel(parent)
        popup.title("O programie")
        popup.resizable(False, False)
        popup.grab_set()
        
        # import danych o wersji programu z pliku version.py
        from project.config.version import (
            PROGRAM_NAME,
            PROGRAM_VERSION,
            PROGRAM_YEAR,
            PROGRAM_AUTHOR,
            DESCRIPTION,
            COMPANY_MAIL,
            PRIVATE_MAIL
            
        )        

        text = (
            f"{PROGRAM_NAME} \n\n"
            f"{DESCRIPTION}"
            f"Email firmowy: {COMPANY_MAIL}\n"
            f"Email prywatny: {PRIVATE_MAIL}\n\n"
            f"Wersja: {PROGRAM_VERSION}\n\n"
            f"© Rok: {PROGRAM_YEAR} {PROGRAM_AUTHOR} "
        )

        label = customtkinter.CTkLabel(
            popup,
            text=text,
            justify="center",
            wraplength=360,
            font=customtkinter.CTkFont(size=13)
        )
        label.pack(padx=20, pady=20)

        customtkinter.CTkButton(
            popup,
            text="OK",
            command=popup.destroy
        ).pack(pady=(0, 15))

        try:
            center_popup(parent, popup)
        except Exception:
            pass
        
    about_btn = customtkinter.CTkButton(
    left,
    text="O programie",
    command=lambda: show_about_popup(root)
    )
    about_btn.grid(row=99, column=0, pady=(0, 10), sticky="ew")



    # funkcja czyszczenia textboxa
    def clean_text():
        text.configure(state="normal")
        text.delete("1.0", "end")
        text.configure(state="disabled")
        result_var.set("")
        app_state["last_report_kind"] = None
        app_state["last_report_data"] = None
        app_state["last_report_text"] = ""
        _set_print_visible(False)
        _upadate_placeholder_visibility()



# 6) Przykładowe przyciski w lewym panelu
    download_machine_btn = customtkinter.CTkButton(left, text="Wczytaj maszyny", command=lambda: loading_machine_data(root))
    download_machine_btn.grid(row=0, column=0, pady=(0, 10), sticky="ew")
    load_machine_btn = customtkinter.CTkButton(left, text="Wczytaj plik", command=on_open_file)
    load_machine_btn.grid(row=1, column=0, pady=(0, 10), sticky="ew")
    count_production_btn = customtkinter.CTkButton(left, text="Przelicz produkcję", command=calculate_production)
    count_production_btn.grid(row=2, column=0, pady=(0, 10), sticky="ew")
    generate_report_btn = customtkinter.CTkButton(left, text="Generuj raport", command=generate_logistics_report)
    generate_report_btn.grid(row=3, column=0, pady=(0, 10), sticky="ew")

    clean_btn = customtkinter.CTkButton(left, text="Wyczyść", command=clean_text)
    clean_btn.grid(row=4, column=0, pady=(0, 10), sticky="ew")
    # umieść przycisk przełączania motywu w lewym panelu, aby pasował do pozostałych kontrolek
    # ustaw początkowy tekst zgodnie z aktualnym trybem wyglądu

    toogle_text = "Jasny motyw" if customtkinter.get_appearance_mode() == "Dark" else "Ciemny motyw"
    ch_theme = customtkinter.CTkButton(left, text=toogle_text, command=change)
    ch_theme.grid(row=5, column=0, pady=(20, 10), sticky="ew")

    root.mainloop()